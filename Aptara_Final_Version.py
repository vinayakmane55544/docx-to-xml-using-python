import docx
from xml.etree.ElementTree import Element, SubElement, ElementTree
import xml.etree.ElementTree as ET

def remove_superscripts(paragraphs):
    non_superscripts = []

    for paragraph in paragraphs:
        for run in paragraph.runs:
            # Check if the run has a superscript format
            if not run.font or not run.font.superscript:
                non_superscripts.append(run.text)

    return ''.join(non_superscripts)

def extract_headings_and_paragraphs(docx_path):
    doc = docx.Document(docx_path)
    headings_and_paragraphs = []
    current_heading = None
    current_paragraphs = []
    use_is_heading2 = False

    for paragraph in doc.paragraphs:
        if use_is_heading2:
            if is_heading(paragraph):
                if current_heading and current_paragraphs:
                    headings_and_paragraphs.append((current_heading, current_paragraphs))
                current_heading = paragraph.text
                current_paragraphs = []
            else:
                current_paragraphs.append(paragraph)
        else:
            if is_heading2(paragraph):
                if current_heading and current_paragraphs:
                    headings_and_paragraphs.append((current_heading, current_paragraphs))
                current_heading = paragraph.text
                current_paragraphs = []
            else:
                current_paragraphs.append(paragraph)
                # If is_heading didn't match, try is_heading2
                if current_heading is None:
                    use_is_heading2 = True

    if current_heading and current_paragraphs:
        headings_and_paragraphs.append((current_heading, current_paragraphs))

    return headings_and_paragraphs

def is_heading(paragraph):
    return "Heading" in paragraph.style.name

def is_heading2(paragraph):
    return any(run.bold for run in paragraph.runs)

def removeSpaces(string):
    return string.replace(' ', '')

def main():
    docx_path = '4.docx'  # Provide the correct input Word document path

    # Extract title from DOCX
    docx_document = docx.Document(docx_path)
    title = docx_document.paragraphs[0].text

    # Create the XML structure
    root = Element("root")
    tree = ElementTree(root)

    # Add title to XML
    title_group = SubElement(root, "titleGroup")
    title_element = SubElement(title_group, "title", {"type": "main"})
    title_element.text = title

    # Continue with the rest of the XML creation logic
    headings_and_paragraphs = extract_headings_and_paragraphs(docx_path)

    if headings_and_paragraphs:
        for heading, paragraphs in headings_and_paragraphs:
            print(heading)
            if removeSpaces(heading) == "Keywords":
                keywords_root = SubElement(root, "keywordGroup", type="author")

                for paragraph in paragraphs:
                    keywords = [keyword.strip() for keyword in paragraph.text.split(',')]
                    for index, keyword in enumerate(keywords):
                        keyword_element = SubElement(keywords_root, "keyword",
                                                     xml_id=f"aenm202204208-kwd-{index + 1:04d}")
                        keyword_element.text = keyword
                        keyword_element.tail = "\n"

            elif heading == "Abstract":
                abstract_root = SubElement(root, "abstract", type="main", xml_lang="en")
                title = SubElement(abstract_root, "title", type="main")
                title.text = "Abstract"
                for paragraph in paragraphs:
                    p = SubElement(abstract_root, "p")
                    p.text = paragraph.text

            elif heading == "Introduction":
                introduction_root = SubElement(root, "introduction", type="main")
                title = SubElement(introduction_root, "title", type="main")
                title.text = "Introduction"
                for paragraph in paragraphs:
                    p = SubElement(introduction_root, "p")
                    p.text = paragraph.text

            elif removeSpaces(heading) == "Authors":
                authors_root = SubElement(root, "creators")
                authors_root.tail = "\n"

                for paragraph in paragraphs:
                    authors_p = remove_superscripts([paragraph])
                    # Split author names and create XML structure
                    authors = [author.strip() for author in authors_p.split(',') if len(author) > 2]
                    for index, author in enumerate(authors):
                        names = author.split()
                        given_names = ' '.join(names[:-1])
                        family_name = names[-1]

                        creator = SubElement(authors_root, "creator", xml_id=f"aenm202204208-cr-{index + 1:04d}",
                                             creatorRole="author")
                        creator.tail = "\n"
                        person_name = SubElement(creator, "personName")
                        given_names_elem = SubElement(person_name, "givenNames")
                        given_names_elem.text = given_names
                        family_name_elem = SubElement(person_name, "familyName")
                        family_name_elem.text = family_name

                    contact_details = SubElement(creator, "contactDetails")
                    email_elem = SubElement(contact_details, "email")
                    email_elem.text = "ref@gmail.com"


            elif removeSpaces(heading) == "Affiliation":
                affiliation_group = SubElement(root, "affiliationGroup")
                affiliation_group.tail = "\n"

                for paragraph in paragraphs:
                    affiliation_info = remove_superscripts([paragraph])
                    affiliation_info = [affiliation_info.strip()]
                    print(affiliation_info)
                    for index, affiliation_info in enumerate(affiliation_info):

                        affiliation = SubElement(affiliation_group, "affiliation", xml_id=f"aenm202204208-aff-{index + 1:04d}",
                                                 countryCode="US")
                        affiliation.tail = "\n"

                        org_div_elem = SubElement(affiliation, "orgDiv")
                        org_div_elem.text = "custom input"
                        org_name_elem = SubElement(affiliation, "orgName")
                        org_name_elem.text = affiliation_info

                        address = SubElement(affiliation_group, "address")
                        street_elem = SubElement(address, "street")
                        street_elem.text = "custom input"
                        city_elem = SubElement(address, "city")
                        city_elem.text = "custom input"
                        country_part_elem = SubElement(address, "countryPart")
                        country_part_elem.text = "custom input"
                        post_code_elem = SubElement(address, "postCode")
                        post_code_elem.text = "custom input"
                        country_elem = SubElement(address, "country")
                        country_elem.text = "custom input"

    # Write the entire XML structure to the XML file
    output_xml_path = 'v4.4_output.xml'  # Provide the correct output XML file path
    tree.write(output_xml_path, encoding="utf-8", xml_declaration=True)
    print(f"XML file generated successfully at {output_xml_path}.")

if __name__ == "__main__":
    main()
