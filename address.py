import docx
from xml.etree.ElementTree import Element, SubElement, ElementTree

def extract_headings_and_paragraphs(docx_path):
    doc = docx.Document(docx_path)
    headings_and_paragraphs = []
    current_heading = None
    current_paragraph = []
    use_is_heading2 = False

    for paragraph in doc.paragraphs:
        if use_is_heading2:
            if is_heading(paragraph):
                if current_heading and current_paragraph:
                    headings_and_paragraphs.append((current_heading, '\n'.join(current_paragraph)))
                current_heading = paragraph.text
                current_paragraph = []
            else:
                current_paragraph.append(paragraph.text)
        else:
            if is_heading2(paragraph):
                if current_heading and current_paragraph:
                    headings_and_paragraphs.append((current_heading, '\n'.join(current_paragraph)))
                current_heading = paragraph.text
                current_paragraph = []
            else:
                current_paragraph.append(paragraph.text)
                # If is_heading didn't match, try is_heading2
                if current_heading is None:
                    use_is_heading2 = True

    if current_heading and current_paragraph:
        headings_and_paragraphs.append((current_heading, '\n'.join(current_paragraph)))

    return headings_and_paragraphs

def is_heading(paragraph):
    return "Heading" in paragraph.style.name

def is_heading2(paragraph):
    return any(run.bold for run in paragraph.runs)

def removeSpaces(string):
    string = string.replace(' ','')
    return string

def handle_superscripts(paragraph, parent_element):
    # Iterate through runs in the paragraph to look for superscripts
    for run in paragraph.runs:
        if run.font.superscript:
            # Create a new <sup> element for the superscripted text
            sup_element = SubElement(parent_element, "sup")
            sup_element.text = run.text
        else:
            # Create a regular <text> element for non-superscripted text
            text_element = SubElement(parent_element, "text")
            text_element.text = run.text


def extract_contact_details(docx_path):
    doc = docx.Document(docx_path)
    in_contact_details = False
    address_info = {}
    expected_keys = ["street", "city", "countryPart", "postCode", "country"]
    current_key_index = 0

    for paragraph in doc.paragraphs:
        if "Contact details" in paragraph.text:
            in_contact_details = True
        elif in_contact_details:
            line = paragraph.text.strip()
            if current_key_index < len(expected_keys):
                key = expected_keys[current_key_index]
                address_info[key] = line
                current_key_index += 1

    return address_info

def create_address_xml(address_info):
    root = Element("root")
    address_root = SubElement(root, "address")

    for key, value in address_info.items():
        element = SubElement(address_root, key)
        element.text = value

    return root

# Example usage
docx_path = '1.docx'

headings_and_paragraphs = extract_headings_and_paragraphs(docx_path)
contact_details_info = extract_contact_details(docx_path)

root = Element("root")
tree = ElementTree(root)

if headings_and_paragraphs:
    for heading, paragraph in headings_and_paragraphs:
        print(heading)
        if removeSpaces(heading) == "Keywords":
            # print(f'Paragraph:\n{paragraph}\n')
            keywords = [keyword.strip() for keyword in paragraph.split(',')]
            keywords_root = SubElement(root, "keywordGroup", type="author")
            # print(keywords)

            for index, keyword in enumerate(keywords):
                keyword_element = SubElement(keywords_root, "keyword", xml_id=f"aenm202204208-kwd-{index + 1:04d}")
                keyword_element.text = keyword
                keyword_element.tail = "\n"

        elif heading == "Abstract":
            abstract_root = SubElement(root, "abstract", type="main", xml_lang="en")
            title = SubElement(abstract_root, "title", type="main")
            title.text = "Abstract"
            p = SubElement(abstract_root, "p")
            p.text = paragraph

        elif heading == "Introduction":
            introduction_root = SubElement(root, "introduction", type="main")
            title = SubElement(introduction_root, "title", type="main")
            title.text = "Introduction"
            p = SubElement(introduction_root, "p")
            p.text = paragraph

        elif removeSpaces(heading) == "Authors":
            authors_root = SubElement(root, "creators")

            # Split author names and create XML structure
            authors = [author.strip() for author in paragraph.split(',')]
            for index, author in enumerate(authors):
                names = author.split()
                given_names = ' '.join(names[:-1])
                family_name = names[-1]

                creator = SubElement(authors_root, "creator", xml_id=f"aenm202204208-cr-{index + 1:04d}",
                                     creatorRole="author")
                person_name = SubElement(creator, "personName")
                given_names_elem = SubElement(person_name, "givenNames")
                given_names_elem.text = given_names
                family_name_elem = SubElement(person_name, "familyName")
                family_name_elem.text = family_name

if contact_details_info:
    address_xml = create_address_xml(contact_details_info)

    # Add the address XML to the existing structure
    root.append(address_xml)

# Write the entire XML structure to the XML file
tree.write("v1.01_output_combined.xml", encoding="utf-8", xml_declaration=True)
print("XML file generated successfully.")